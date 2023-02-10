"""
Web Scrapper made by : Naman Garg
Section : CS A 04
Roll Number : 12112061
Contact: <12112061@nitkkr.ac.in>

NIT Scrapped in this project are:

1. NIT Uttarakhand
2. NIT Puducherry
3. NIT Arunachal Pradesh
4. NIT Sikkim
5. NIT Delhi
6. NIT Mizoram
7. NIT Nagaland
8. NIT Andhra Pradesh
"""

from selenium import webdriver
from bs4 import BeautifulSoup
from selenium.webdriver.firefox.options import Options

import docx

doc = docx.Document()

options = Options()
options.add_argument("--headless")
driver = webdriver.Firefox(options=options)

#============================================================================================= 1. NIT Uttarakhand =============================================================================================

doc.add_heading("1. NIT Uttarakhand", 0)
doc.add_paragraph("Full Name \n Designation \n Email ID \n Contact Number \n Research area")

url_list = ['https://nituk.ac.in/computer-science-engineering/peoples', 'https://nituk.ac.in/civil-engineering/peoples',
            'https://nituk.ac.in/electrical-engineering/peoples', 'https://nituk.ac.in/electronics-engineering/peoples',
            'https://nituk.ac.in/mechanical-engineering/peoples'
            ]
for page in url_list:
    driver.get(page)
    soup = BeautifulSoup(driver.page_source, "html.parser")
    s = soup.find_all("td")

    for faculty in s:
        facultyData = faculty.text

        # Skip, because not all td are faculty
        if facultyData.find("Designation") == -1:
            continue

        facultyData = facultyData.strip().split("\n")
        # Faculty Name
        fac_name = facultyData[0].strip()

        # Designation
        fac_designation = facultyData[2].split(":")[1]

        # Email ID
        if len(facultyData) >= 5:
            fac_email = facultyData[4].split(":")[1].strip()
        else:
            fac_email = "-"

        fac_phone = "-"
        # Phone Number
        if len(facultyData) >= 4:
            for string in facultyData[3].split():
                if string.isnumeric():
                    fac_phone = string
                    break

        fac_research = "-"

        if len(facultyData) >= 6:
            fac_research = facultyData[5].split(":")[1].strip()

        doc.add_paragraph("\n==============================================================\n" +
                          fac_name + " \n" + fac_designation + " \n " + fac_email + " \n " +
                          fac_phone + " \n " + fac_research)

doc.add_page_break()
doc.save('output.docx')
print("Done NIT Uttarakhand")

# ============================================================================================= 2. NIT Puducherry =============================================================================================

doc.add_heading("2. NIT Puducherry", 0)
doc.add_paragraph("Full Name \n Designation \n Email ID \n Contact Number \n Research area")

url_list = ["https://nitpy.ac.in/academics/departments/civil/faculty",
            "https://nitpy.ac.in/academics/departments/cse/faculty",
            "https://nitpy.ac.in/academics/departments/ece/faculty",
            "https://nitpy.ac.in/academics/departments/eee/faculty",
            "https://nitpy.ac.in/academics/departments/mech/faculty"
            ]

for page in url_list:
    driver.get(page)
    driver.refresh()
    soup = BeautifulSoup(driver.page_source, "html.parser")
    s = soup.find_all("app-department-faculty-view")
    for faculty in s:
        facultyData = faculty.find_all("p")

        # Faculty Name
        fac_name = facultyData[0].text.split(":")[1].strip()

        # Designation
        fac_designation = facultyData[1].text.split(":")[1].strip()

        # E-Mail
        fac_email = facultyData[2].text.split(":")[1].strip()

        # Phone Number
        fac_phone = facultyData[3].text.split(":")[1].strip()

        fac_research = "-"

        if len(facultyData) >= 5:
            try:
                fac_research = facultyData[4].text.split(":")[1].strip()
            except:
                fac_research = "-"
            finally:
                pass
        doc.add_paragraph("\n==============================================================\n" +
                          fac_name + " \n" + fac_designation + " \n " + fac_email + " \n " +
                          fac_phone + " \n " + fac_research)

doc.save('output.docx')
doc.add_page_break()
print("Done NIT Puducherry")

# ============================================================================================= 3. NIT Arunachal Pradesh =============================================================================================

doc.add_heading("3. NIT Arunachal Pradesh", 0)
doc.add_paragraph("Full Name \n Designation \n Email ID \n Contact Number \n Research area")

url_list = ["https://www.nitap.ac.in/department/faculty?faculty=33b1f5929e",
            "https://www.nitap.ac.in/department/faculty?faculty=1dbf504017",
            "https://www.nitap.ac.in/department/faculty?faculty=8345cde791",
            "https://www.nitap.ac.in/department/faculty?faculty=66badff20c",
            "https://www.nitap.ac.in/department/faculty?faculty=cad3da575a",
            "https://www.nitap.ac.in/department/faculty?faculty=89f0915482"
            ]

for page in url_list:
    driver.get(page)
    soup = BeautifulSoup(driver.page_source, "html.parser")
    s = soup.find_all("div", class_="gdlr-core-personnel-list-content-wrap")
    for faculty in s:

        # Faculty Name
        fac_name = faculty.find_all("h3")[0].text.strip()

        # Designation
        fac_designation = faculty.find_all("div", class_="gdlr-core-personnel-list-position gdlr-core-info-font gdlr-core-skin-caption")[0].text.strip()

        # E-Mail
        fac_email = faculty.find_all("a")[0].text.strip()

        # Phone Number
        fac_phone = faculty.find_all("div", class_="kingster-personnel-info-list kingster-type-phone")[0].text.strip()

        # Research Area
        fac_research = faculty.find_all("p")[-1].text.strip()

        doc.add_paragraph("\n==============================================================\n" +
                          fac_name + " \n" + fac_designation + " \n " + fac_email + " \n " +
                          fac_phone + " \n " + fac_research)

doc.save('output.docx')
doc.add_page_break()
print("Done NIT Arunachal Pradesh")


# ============================================================================================= 4. NIT Sikkim =============================================================================================

doc.add_heading("4. NIT Sikkim", 0)
doc.add_paragraph("Full Name \n Designation \n Email ID \n Contact Number \n Research area")

url_list = ["https://cse.nitsikkim.ac.in/people.php",
            "https://ece.nitsikkim.ac.in/people.php",
            "https://eee.nitsikkim.ac.in/people.php",
            "https://me.nitsikkim.ac.in/people.php",
            "https://ce.nitsikkim.ac.in/people.php",
            ]

for page in url_list:
    driver.get(page)
    soup = BeautifulSoup(driver.page_source, "html.parser")
    s = soup.find_all("div", class_="card-body")
    for faculty in s:

        # Faculty Name
        fac_name = faculty.find_all("h5")[0].text.strip()

        # Designation
        fac_designation = faculty.find_all("h6")[0].text.strip()

        # E-Mail
        fac_email = faculty.find_all("p", class_="card-text")[0].text.split("+")[0].strip()

        # # Phone Number
        fac_phone = "+" + faculty.find_all("p", class_="card-text")[0].text.split("+")[-1].strip()

        # Research Area
        fac_research = faculty.find_all("p", class_="card-text")[-1].text.strip()
        doc.add_paragraph("\n==============================================================\n" +
                          fac_name + " \n" + fac_designation + " \n " + fac_email + " \n " +
                          fac_phone + " \n " + fac_research)

doc.save('output.docx')
doc.add_page_break()
print("Done NIT Sikkim")

# ============================================================================================= 5. NIT Delhi =============================================================================================

doc.add_heading("5. NIT Delhi", 0)
doc.add_paragraph("Full Name \n Designation \n Email ID \n Contact Number \n Research area")

url_list = ["https://nitdelhi.ac.in/?page_id=11979",
            "https://nitdelhi.ac.in/?page_id=11977",
            "https://nitdelhi.ac.in/?page_id=11985",
            "https://nitdelhi.ac.in/?page_id=11981",
            "https://nitdelhi.ac.in/?page_id=11993"
            ]

for page in url_list:
    driver.get(page)
    soup = BeautifulSoup(driver.page_source, "html.parser")
    s = soup.find_all("tr")
    for faculty in s:

        if len(faculty.find_all("a")) == 0:
            continue
        doc.add_paragraph("\n==============================================================\n" +
                          faculty.find_all("td")[1].text)

doc.save('output.docx')
doc.add_page_break()
print("Done NIT Delhi")

#============================================================================================= 6. NIT Mizoram =============================================================================================

doc.add_heading("6. NIT Mizoram", 0)
doc.add_paragraph("Full Name \n Designation \n Email ID \n Contact Number \n Research area")

url_list = ["https://www.nitmz.ac.in/Department_FaculyList.aspx?nDeptID=cg",
            "https://www.nitmz.ac.in/Department_FaculyList.aspx?nDeptID=ec",
            "https://www.nitmz.ac.in/Department_FaculyList.aspx?nDeptID=ck",
            "https://www.nitmz.ac.in/Department_FaculyList.aspx?nDeptID=ci",
            "https://www.nitmz.ac.in/Department_FaculyList.aspx?nDeptID=cm"
            ]

for page in url_list:
    driver.get(page)
    soup = BeautifulSoup(driver.page_source, "html.parser")
    s = soup.find_all("table", class_="DataGrid")[0].find_all("tr")
    for faculty in s:
        if len(faculty.find_all("a")) == 0:
            continue
        # Faculty Name
        fac_name = faculty.find_all("a")[0].text.strip()

        # Designation
        fac_designation = faculty.find_all("td")[1].text.strip()

        # E-Mail
        fac_email = faculty.find_all("td")[2].text.strip()

        doc.add_paragraph("\n==============================================================\n" +
                          fac_name + " \n" + fac_designation + " \n " + fac_email + " \n " +
                          "-" + " \n " + "-")

doc.save('output.docx')
doc.add_page_break()
print("Done NIT Mizoram")

#============================================================================================= 7. NIT Nagaland =============================================================================================

doc.add_heading("7. NIT Nagaland", 0)
doc.add_paragraph("Full Name \n Designation \n Email ID \n Contact Number \n Research area")

url_list = ["https://www.nitnagaland.ac.in/index.php/cse-people/cse-faculty",
            "https://www.nitnagaland.ac.in/index.php/eee-people/eee-faculty",
            "https://www.nitnagaland.ac.in/index.php/ece-people/ece-faculty",
            "https://www.nitnagaland.ac.in/index.php/eie-people/eie-faculty",
            ]

for page in url_list:
    driver.get(page)
    soup = BeautifulSoup(driver.page_source, "html.parser")
    s = soup.find_all("div", class_="fp-testimonials")
    for faculty in s:
        if len(faculty.find_all("b")) <= 1:
            continue

        # Faculty Name
        fac_name = faculty.find_all("b")[0].text.strip()
        # Designation
        fac_designation = faculty.find_all("b")[1].text.strip()

        # E-Mail
        fac_email = faculty.text.split("\n")[5].strip()

        # Research Area
        fac_research = faculty.text.split("\n")[6].strip()
        doc.add_paragraph("\n==============================================================\n" +
                          fac_name + " \n" + fac_designation + " \n " + fac_email + " \n " +
                          "-" + " \n " + fac_research)

doc.save('output.docx')
doc.add_page_break()
print("Done NIT Nagaland")
